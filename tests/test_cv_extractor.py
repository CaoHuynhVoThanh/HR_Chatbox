import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document

from app.core.cv_extractor import CVExtractionError, extract_cv, resolve_cv_path


class CVExtractorTests(unittest.TestCase):
    def test_extracts_docx_paragraphs_and_tables(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            cv_path = Path(tmp_dir) / "candidate.docx"
            document = Document()
            document.add_paragraph("Nguyen Van A - Python Developer")
            table = document.add_table(rows=1, cols=1)
            table.cell(0, 0).text = "FastAPI"
            document.save(cv_path)

            result = extract_cv(cv_path)

        self.assertIn("Python Developer", result.text)
        self.assertIn("FastAPI", result.text)

    def test_requires_explicit_path_when_cv_directory_has_multiple_files(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            directory = Path(tmp_dir)
            (directory / "a.pdf").touch()
            (directory / "b.docx").touch()
            with self.assertRaises(CVExtractionError):
                resolve_cv_path(None, directory)

    @patch("app.core.cv_extractor._extract_pdf_with_pymupdf")
    @patch("app.core.cv_extractor._extract_pdf_with_pypdf")
    def test_pdf_uses_cleaner_fallback_extractor(self, pypdf_extract, pymupdf_extract) -> None:
        pypdf_extract.return_value = "CAO HU\u25a0NH V\u00d5 THANH"
        pymupdf_extract.return_value = "CAO HU\u1ef2NH V\u00d5 THANH"

        with tempfile.TemporaryDirectory() as tmp_dir:
            cv_path = Path(tmp_dir) / "candidate.pdf"
            cv_path.touch()
            result = extract_cv(cv_path)

        self.assertEqual(result.text, "CAO HU\u1ef2NH V\u00d5 THANH")

    @patch("app.core.cv_extractor._extract_pdf_with_pymupdf")
    @patch("app.core.cv_extractor._extract_pdf_with_pypdf")
    def test_pdf_with_unmapped_unicode_fails_clearly(self, pypdf_extract, pymupdf_extract) -> None:
        corrupted_text = "Cao Hu\u25a0nh V\u25a0 Thanh " + ("\u25a0" * 20)
        pypdf_extract.return_value = corrupted_text
        pymupdf_extract.return_value = corrupted_text

        with tempfile.TemporaryDirectory() as tmp_dir:
            cv_path = Path(tmp_dir) / "candidate.pdf"
            cv_path.touch()
            with self.assertRaisesRegex(CVExtractionError, "không ánh xạ được Unicode"):
                extract_cv(cv_path)
