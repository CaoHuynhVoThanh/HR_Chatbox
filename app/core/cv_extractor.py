from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
import pymupdf
from pypdf import PdfReader

SUPPORTED_EXTENSIONS = {".pdf", ".docx"}

ICON_MAP = {
    "\uf095": "Phone:",
    "\uf0e0": "Email:",
    "\uf015": "Address:",
    "\uf005": "★",
    "\uf0ac": "Website:",
    "\uf09b": "GitHub:",
    "\uf08c": "LinkedIn:",
    "\uf007": "Name:",
    "\uf0b1": "Work:",
    "\uf19d": "Education:",
    "\uf133": "Date:",
    "\uf041": "Location:",
    "\uf0c0": "Team:",
    "\uf0ae": "Skills:",
    "\uf0c3": "Phone:",
}


def normalize_cv_text(text: str) -> str:
    normalized = text
    for unicode_char, replacement in ICON_MAP.items():
        normalized = normalized.replace(unicode_char, replacement)
    return normalized


def _extract_pdf_with_pypdf(path: Path) -> str:
    return "\n".join(page.extract_text() or "" for page in PdfReader(str(path)).pages)


def _extract_pdf_with_pymupdf(path: Path) -> str:
    """Use PyMuPDF as a second text-extraction engine for PDF font encodings."""
    with pymupdf.open(path) as document:
        return "\n".join(page.get_text("text", sort=True) for page in document)


def _pdf_text_quality(text: str) -> tuple[float, int]:
    """Return a score where lower replacement-character rate is better."""
    visible_characters = sum(not char.isspace() for char in text)
    unreadable_characters = sum(char in {"\u25a0", "\ufffd", "\x00"} for char in text)
    return (unreadable_characters / max(visible_characters, 1), -visible_characters)


def _has_unrecoverable_pdf_text(text: str) -> bool:
    """Detect the common missing-ToUnicode symptom without rejecting normal bullets."""
    unreadable_characters = sum(char in {"\u25a0", "\ufffd", "\x00"} for char in text)
    visible_characters = sum(not char.isspace() for char in text)
    return unreadable_characters >= 4 and unreadable_characters / max(visible_characters, 1) >= 0.01


def _extract_pdf_text(path: Path) -> str:
    """Choose the cleanest result from two native PDF text extractors.

    Neither extractor can recreate letters that were never mapped to Unicode in
    the PDF. In that case, fail explicitly so corrupted CV content is never sent
    to Gemini as if it were valid Vietnamese text.
    """
    candidates: list[str] = []
    errors: list[Exception] = []
    for extractor in (_extract_pdf_with_pypdf, _extract_pdf_with_pymupdf):
        try:
            extracted = extractor(path).strip()
            if extracted:
                candidates.append(extracted)
        except Exception as exc:
            errors.append(exc)

    if not candidates:
        if errors:
            raise errors[0]
        return ""

    text = min(candidates, key=_pdf_text_quality)
    if _has_unrecoverable_pdf_text(text):
        raise CVExtractionError(
            "PDF dùng font không ánh xạ được Unicode nên văn bản tiếng Việt bị lỗi. "
            "Hãy tải CV .docx hoặc xuất lại PDF có thể chọn/copy được chữ; PDF scan cần OCR."
        )
    return text


class CVExtractionError(ValueError):
    """Raised when a CV cannot be selected or converted into text."""


@dataclass(frozen=True)
class CVDocument:
    path: Path
    text: str


def resolve_cv_path(cv_path: str | Path | None, cv_directory: Path = Path("data/cv")) -> Path:
    """Find a CV. The directory selection is intentionally non-strict for this demo.

    A future upload API can enforce exactly one stored CV per user. Until then,
    passing --cv is the least ambiguous option when the folder contains many files.
    """
    if cv_path is not None:
        path = Path(cv_path)
        if not path.is_file():
            raise CVExtractionError(f"Không tìm thấy CV: {path}")
        return path

    files = sorted(
        path for path in cv_directory.glob("*") if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS
    )
    if not files:
        raise CVExtractionError(
            f"Chưa có CV .pdf/.docx trong '{cv_directory}'. Hãy thêm CV hoặc dùng --cv <đường_dẫn>."
        )
    if len(files) > 1:
        choices = ", ".join(path.name for path in files)
        raise CVExtractionError(f"Có nhiều CV: {choices}. Hãy chỉ định file bằng --cv.")
    return files[0]


def extract_cv(path: str | Path) -> CVDocument:
    cv_path = Path(path)
    extension = cv_path.suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise CVExtractionError("Chỉ hỗ trợ tệp PDF (.pdf) và Word (.docx).")

    try:
        if extension == ".pdf":
            text = _extract_pdf_text(cv_path)
        else:
            document = Document(str(cv_path))
            paragraphs = [paragraph.text for paragraph in document.paragraphs]
            table_cells = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
            text = "\n".join(part for part in [*paragraphs, *table_cells] if part.strip())
    except Exception as exc:  # Libraries expose several format-specific exceptions.
        raise CVExtractionError(f"Không thể đọc CV '{cv_path.name}': {exc}") from exc

    text = text.strip()
    if not text:
        raise CVExtractionError(
            "Không trích xuất được chữ từ CV. Với PDF scan, cần bổ sung OCR ở giai đoạn sau."
        )
    return CVDocument(path=cv_path, text=normalize_cv_text(text))
